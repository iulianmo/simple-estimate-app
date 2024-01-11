#!/usr/bin/env python3.10

import tkinter as tk
from tkinter import ttk 
from tkinter.ttk import Combobox
from data import dati
from docx import Document 
#from docx.shared import Pt       
#from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import messagebox
import pathlib


k = dati.keys()                                                                            # inserisce nella variabile k i nomi dei servizi
v = dati.values()                                                                          # inserisce nella variabile v i prezzi

# ----- LISTE  -----
lista_y = [200]                                                                             # primo valore della lista y
lista_x = [50]                                                                              # primo valore della lista x
lista_servizi = []
lista_prezzo = []
lista_misura = []
lista_n = [0]


# ----- FUN. AUTOCOMPLETAMENTO  -----
class AutocompleteCombobox(ttk.Combobox):                                                  # funzione per autocompletamento
        def set_completion_list(self, completion_list):
                self._completion_list = list(k)                                            # Work with a sorted list
                self._hits = []
                self._hit_index = 0
                self.position = 0
                self.bind('<KeyRelease>', self.handle_keyrelease)
                self['values'] = self._completion_list                                     # Setup our popup menu

        def autocomplete(self, delta=0):
                if delta:                                                                  # need to delete selection otherwise we would fix the current position
                        self.delete(self.position, tk.END)
                else:                                                                      # set position to end so selection starts where textentry ended
                        self.position = len(self.get())
                _hits = []                                                                 # collect hits
                for element in self._completion_list:
                        if element.lower().startswith(self.get().lower()):                 # Match case insensitively
                                _hits.append(element)
                if _hits != self._hits:                                                    # if we have a new hit list, keep this in mind
                        self._hit_index = 0
                        self._hits=_hits
                if _hits == self._hits and self._hits:                                     # only allow cycling if we are in a known hit list
                        self._hit_index = (self._hit_index + delta) % len(self._hits)
                if self._hits:                                                             # now finally perform the auto completion
                        self.delete(0,tk.END)
                        self.insert(0,self._hits[self._hit_index])
                        self.select_range(self.position,tk.END)

        def handle_keyrelease(self, event):
                if event.keysym == "BackSpace":
                        self.delete(self.index(tk.INSERT), tk.END)
                        self.position = self.index(tk.END)
                if event.keysym == "Left":
                        if self.position < self.index(tk.END):                              # delete the selection
                                self.delete(self.position, tk.END)
                        else:
                                self.position = self.position-1                             # delete one character
                                self.delete(self.position, tk.END)
                if event.keysym == "Right":
                        self.position = self.index(tk.END)                                  # go to end (no selection)
                if len(event.keysym) == 1:
                        self.autocomplete()


# ----- FUN. DOMANDA DI USCITA APP -----
def ask_exit():                                                                            # chiede prima di uscire dall'app       
    if messagebox.askokcancel("Exit", "Vuoi uscire dall'applicazione?"):
        app.destroy()


# ----- FUN. AGGIUNTA SERVIZIO  -----
def add_servizio():                                                                          # funzione aggiunta nuovo servizio e prezzo
        valore_x = lista_x[0]                                                                # legge il primo valore della lista
        valore_y = lista_y[0]                                                                # legge il primo valore della lista

        if valore_y == 800:                                                                  # serve per portare Combobox alla seconda parte
                valore_y = 200
                valore_x += 700

                lista_y.insert(0, valore_y) 
                lista_x.insert(0,valore_x)

        servizio = AutocompleteCombobox(app, height=30, width=40)                            # definisce il menu a tendina con il autocompletamento
        servizio.set_completion_list(k)                                                      # definisce i valori del menu
        servizio.place(x=valore_x, y=valore_y)                                               # posizione menu  
        servizio.focus_set()

        prezzo = tk.Entry(app, width=7, fg='black', bg='white')                                                     # definisce l'entry del prezzo
        prezzo.place(x=int(valore_x + 420), y=valore_y)                                      # posizione entry (riprende il valore di x e aggiunge 450)
        prezzo.focus_set()

        misura = tk.Entry(app, width=7, fg='black', bg='white')                                                     # definisce l'entry della misura
        misura.place(x=int(valore_x + 530), y=valore_y)                                      # posizione entry (riprende il valore di x e aggiunge 450)
        misura.focus_set()

        valore_y += 50                                                                       # incrementa la vaiabile per la posizione del menu 
        lista_y.insert(0, valore_y)                                                          # inserisce il nuovo valore di y nella lista

        lista_servizi.append(servizio)                                                       # inserisce nella lista il Combobox appena aggiunto
        lista_prezzo.append(prezzo)
        lista_misura.append(misura)

        n = lista_n[-1]                                                                      # il numero di servizi ogni volta incrementa, legge l'ultimo valore inserito e lo incrementa
        n += 1
        lista_n.append(n)


        # ----- FUN. TOTALE -----
        def totale_prezzo():
                lista_totale = []
                i = 0 

                while i < int(len(lista_n) -1):                                                                 # crea un loop per far leggere tutti i prezzi, int(len(i) -1) serve per evitare l'errore "IndexError: list index out of range"
                        valore_prezzo = float(lista_prezzo[i].get())
                        valore_misura = float(lista_misura[i].get())
                        lista_totale.append(valore_prezzo * valore_misura)                                     # fa la moltiplicazione tra il prezzo e la misura (m)
                        valore_y = 200                                                                         
                        
                        for i in range(len(lista_totale)):                                                      # legge il totale di ciascun servizio e cre un'etichetta 
                                totale_servizio_label = tk.Label(app, text = str(lista_totale[i]) + ' €', font=('arial', 15, 'normal'), bg = 'black')    # etichetta 
                                totale_servizio_label.place(x=valore_x + 630, y= valore_y)
                                valore_y += 50

                        totale = sum(lista_totale)                                                            # fa la somma dei prezzi
                        totale_label = tk.Label(app, text = str(totale) + ' €', font=('arial', 18, 'bold'), bg = 'black')    # etichetta numero totale 
                        totale_label.place(x=500, rely=0.91)  
                        i += 1  

        btn_totale = tk.Button(app, text="Total =",font=('arial', 15, 'normal'), command=totale_prezzo)            # bottone aggiunta nuovo servizio
        btn_totale.place(x=350, rely=0.9, width=125, height=50) 


        # ----- FUN. CREA PREVENTIVO -----
        def create_preventivo():
            nome_preventivo = preventivo.get()

            if nome_preventivo == '':
                messagebox.showwarning("Warning", "Rinonima il tuo preventivo! Il nome del preventivo non può essere vuoto.")
            else: 
                doc = Document()
                titolo = doc.add_heading(level=0) 
                titolo.add_run(nome_preventivo)
                titolo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #stitolo.font.size = Pt(20)
                #titolo.font.color.rgb = RGBColor(0,0,0)

                lista_totale = []
                i = 0                                                                             # variabile per il ciclo per farla incrementare
            
                while i < int(len(lista_n) -1):                                                   # crea un loop per far leggere tutti i servizi, int(len(i) -1) serve per evitare l'errore "IndexError: list index out of range"
                    valore_servizio = lista_servizi[i].get()
                    valore_prezzo = float(lista_prezzo[i].get())
                    valore_misura = float(lista_misura[i].get())
                    lista_totale.append(float(valore_prezzo) * float(valore_misura))                                     # fa la moltiplicazione tra il prezzo e la misura (m)
                    totale = sum(lista_totale)

                    doc.add_paragraph(valore_servizio) 
                    text_right = doc.add_paragraph('= ' + str(valore_prezzo) + ' €')
                    text_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    i += 1                                                                        # incrementa                           
                totale_paragrafo = doc.add_paragraph()
                totale_paragrafo.add_run('\n\n\nTotal').bold = True                                # parola totale in grassetto
                totale_paragrafo_2 = doc.add_paragraph()
                totale_paragrafo_2.add_run('= ' + str(totale) + (' €')).bold = True                          # prezzo totale in grassettp
                totale_paragrafo_2.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
                path = str(pathlib.Path(__file__).parent.absolute())                                # get the path of where is the application
                doc.save(path + '/' +  nome_preventivo + '.docx')
                messagebox.showinfo("Information", "Il tuo preventivo è stato creato in:\n" + path)    

        btn_get = tk.Button(app, text="CREATE ESTIMATE", font=('arial', 18, 'normal'), command=create_preventivo)            # bottone creazione servizio
        btn_get.place(x=1000, rely=0.9, width=400, height=50) 


# ----- WIDGET -----
app = tk.Tk(className='App Preventivi')                                                       # definisce app come widget principale e il nome
# app.iconbitmap('/path/to/ico/icon.ico')
app.geometry('2000x2000')                                                                     # dimensione widget
app.configure(bg='#96CEB4')                                                                    # coloe sfondo
app.resizable()                                                                               # il widget si può ridimensionare

# ----- NOME PREVENTIVO -----
preventivo = tk.StringVar()                                                                   # definisce la variabile come stringa
name_label = tk.Label(app, text = 'Estimate name:', font=('arial', 30, 'bold'), bg = '#96CEB4')   
name_label.place(x=50, y=50)  
name_input = tk.Entry(app, textvariable = preventivo, font=('arial', 30, 'normal'), fg='black', bg='white')  
name_input.place(relx=0.3, y=50, width=700)

# ----- SERVIZIO -----
servizio_label = tk.Label(app, text = 'Service', font=('arial', 20, 'normal'), bg = '#96CEB4')  # etichetta servizio
servizio_label.place(x=50, y=150)                                                               # posizione
btn_servizio = tk.Button(app, text="Add Service", font=('arial', 15, 'normal'), command=add_servizio)                  # bottone aggiunta nuovo servizio
btn_servizio.place(x=50, rely=0.9, width=250, height=50)                                                               # posizione

# ----- PREZZO -----
prezzo_label = tk.Label(app, text = 'Price', font=('arial', 20, 'normal'), bg = '#96CEB4')    # etichetta servizio
prezzo_label.place(x=465, y=150)  

# ----- MISURA -----
misura_label = tk.Label(app, text = 'Measure', font=('arial', 20, 'normal'), bg = '#96CEB4')    # etichetta 
misura_label.place(x=575, y=150)


app.bind('<Control-Q>', lambda event=None: app.destroy())                                     # scorciatoia per uscire
app.bind('<Control-q>', lambda event=None: app.destroy())
app.protocol("WM_DELETE_WINDOW", ask_exit)                                                    # quanto premi "x" per uscire ti fa la domanda

if __name__ == "__main__":
    app.mainloop()